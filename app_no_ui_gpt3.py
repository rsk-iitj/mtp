import os
from docx import Document
from nlp_processing.file_reader import extract_text_from_file_no_ui
from nlp_tasks.keyword_extraction import extract_keywords
from nlp_tasks.sentiment_analysis import analyze_sentiment
from nlp_tasks.summarization import summarize_text
from nlp_tasks.generate_test_plan_gpt3 import generate_test_plan

# Configuration
project_name = 'gmail'
data_directory = f'data/{project_name}'
output_directory = f'output/generated_plan/{project_name}'
os.makedirs(output_directory, exist_ok=True)

# Domain and tech stack configuration
domain = "E-Commerce"
tech_stack = {
    "Frontend Technology": "React",
    "Backend Technology": "Node.js",
    "Database": "MongoDB",
    "Messaging Queue": "Kafka",
    "Cloud Infrastructure": "AWS",
    "Additional Tech": "GraphQL"
}

# Test resources
test_resources = {
    "test_automation_required": "Yes",
    "num_testers": 5,
    "num_automation_testers": 2,
    "num_test_lead": 1,
    "num_test_managers": 1
}

# Reading documents
user_stories_text = []
for filename in os.listdir(data_directory):
    if filename.endswith('.docx') or filename.endswith('.pdf') or filename.endswith('.txt'):
        filepath = os.path.join(data_directory, filename)
        text = extract_text_from_file_no_ui(filepath)
        user_stories_text.append(text)

user_stories_text = "\n\n".join(user_stories_text)
print(user_stories_text)

# NLP Processing
summary = summarize_text(user_stories_text)
keywords = extract_keywords(user_stories_text)
sentiment_label, _ = analyze_sentiment(user_stories_text)
print(summary)
print(keywords)

# Generate Test Plan
test_plan_content = generate_test_plan(summary, keywords, sentiment_label)

# Save to a Word document
doc_path = os.path.join(output_directory, f'{project_name}_Test_Plan.docx')
doc = Document()
doc.add_heading('Test Plan', level=0)
doc.add_paragraph(test_plan_content)
doc.save(doc_path)
print(f'Test plan saved to {doc_path}')
