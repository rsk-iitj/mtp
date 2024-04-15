# file_reader.py
import os

import pdfplumber
from docx import Document
import pandas as pd
import io


def read_pdf(file):
    with pdfplumber.open(file) as pdf:
        pages = [page.extract_text() for page in pdf.pages if page.extract_text() is not None]
    return "\n".join(pages)


def read_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs if para.text])


def read_text(file):
    return file.getvalue().decode('utf-8')


def read_excel(file):
    df = pd.read_excel(file)
    return "\n".join(df.astype(str).apply(lambda x: ' '.join(x.dropna()), axis=1))


def extract_text_from_file(uploaded_file):
    _, ext = os.path.splitext(uploaded_file.name)
    ext = ext.lower()

    if ext == '.pdf':
        return read_pdf(uploaded_file)
    elif ext in ['.docx', '.doc']:  # treating .doc the same as .docx
        return read_docx(uploaded_file)
    elif ext == '.txt':
        return read_text(uploaded_file)
    elif ext in ['.xls', '.xlsx']:
        return read_excel(uploaded_file)
    else:
        return "Unsupported file format"


import os
from docx import Document

def extract_text_from_file_no_ui(filepath):
    text = ""
    _, ext = os.path.splitext(filepath)  # Directly use filepath
    if ext == ".docx":
        doc = Document(filepath)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif ext == ".txt":
        with open(filepath, 'r', encoding='utf-8') as file:
            text = file.read()
    # Add other file type conditions as needed
    return text



